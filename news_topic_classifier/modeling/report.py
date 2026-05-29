from __future__ import annotations

import logging
import os
from pathlib import Path

import matplotlib.pyplot as plt
import mlflow
import numpy as np
import pyarrow.parquet as pq
import seaborn as sns
from google.cloud import storage
from mlflow import MlflowClient
from omegaconf import DictConfig, OmegaConf
from sklearn.metrics import confusion_matrix

logger = logging.getLogger(__name__)


# =============================================================================
# SECTION 1 — Download predictions from GCS
# =============================================================================
def download_predictions(
    gcs_predictions_uri: str,
    gcp_project: str,
    local_dir: str,
) -> Path:
    """
    Download predictions Parquet from GCS to a local directory.

    Parameters
    ----------
    gcs_predictions_uri : str
        Full GCS URI of the predictions Parquet file, e.g.
    gcp_project : str
        GCP project ID for the storage client.
    local_dir : str
        Local directory to write the file into.

    Returns
    -------
    Path
        Local path of the downloaded Parquet file.
    """
    path_no_scheme = gcs_predictions_uri.replace("gs://", "")
    bucket_name = path_no_scheme.split("/")[0]
    blob_name  = "/".join(path_no_scheme.split("/")[1:])

    local_path = Path(local_dir) / "predictions.parquet"
    local_path.parent.mkdir(parents=True, exist_ok=True)

    client = storage.Client(project=gcp_project)
    client.bucket(bucket_name).blob(blob_name).download_to_filename(str(local_path))

    logger.info("Predictions downloaded to %s", local_path)
    return local_path


# =============================================================================
# SECTION 2 — Fetch run data from MLflow
# =============================================================================
def fetch_run_data(run_id: str, tracking_uri: str) -> dict:
    """
    Pull training curves, params, and final metrics for a run from MLflow.

    Parameters
    ----------
    run_id : str
        MLflow run ID to fetch data for.
    tracking_uri : str
        MLflow tracking URI — SQLite locally, GCS in dev/pp/prd.

    Returns
    -------
    dict
        {
            "params":  dict[str, str]   — logged hyperparameters,
            "metrics": dict[str, float] — final metric values,
            "history": dict[str, list[float]] — per-epoch metric values
        }
    """
    mlflow.set_tracking_uri(tracking_uri)
    client = MlflowClient()

    run = client.get_run(run_id)
    params = dict(run.data.params)
    metrics = dict(run.data.metrics)

    history_keys = ["train_loss", "train_acc", "val_loss", "val_acc"]
    history = {
        key: [m.value for m in client.get_metric_history(run_id, key)]
        for key in history_keys
    }

    logger.info("Fetched run %s - %d epochs logged.", run_id, len(history["train_loss"]))

    return {"params": params, "metrics": metrics, "history": history}


# =============================================================================
# SECTION 3 — Plot: training curves
# =============================================================================
def plot_training_curves(history: dict, save_dir: str) -> Path:
    """
    Plot loss and accuracy curves for train and validation sets.

    Parameters
    ----------
    history : dict
        Per-epoch metric lists from fetch_run_data()["history"].
        Expected keys: train_loss, train_acc, val_loss, val_acc.
    save_dir : str
        Local directory to save the figure into.

    Returns
    -------
    Path
        Local path of the saved figure (training_curves.png).
    """
    if not history.get("train_loss"):
        raise ValueError(
            "Training history is empty — make sure run_id points to the training run, "
            "not the predict run."
        )

    epochs = range(1, len(history["train_loss"]) + 1)

    fig, (ax_loss, ax_acc) = plt.subplots(1, 2, figsize=(12, 4))

    # Loss
    ax_loss.plot(epochs, history["train_loss"], label="Train", marker="o")
    ax_loss.plot(epochs, history["val_loss"],   label="Val",   marker="o")
    ax_loss.set_title("Loss")
    ax_loss.set_xlabel("Epoch")
    ax_loss.set_ylabel("Loss")
    ax_loss.legend()

    # Accuracy
    ax_acc.plot(epochs, history["train_acc"], label="Train", marker="o")
    ax_acc.plot(epochs, history["val_acc"],   label="Val",   marker="o")
    ax_acc.set_title("Accuracy")
    ax_acc.set_xlabel("Epoch")
    ax_acc.set_ylabel("Accuracy")
    ax_acc.legend()

    fig.suptitle("Training Curves", fontsize=14, fontweight="bold")
    fig.tight_layout()

    save_path = Path(save_dir) / "training_curves.png"
    save_path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(save_path, dpi=150, bbox_inches="tight")
    plt.close(fig)

    logger.info("Training curves saved to %s", save_path)
    return save_path


# =============================================================================
# SECTION 4 — Plot: confusion matrix
# =============================================================================
def plot_confusion_matrix(
    preds: np.ndarray,
    labels: np.ndarray,
    label_names: list[str],
    save_dir: str,
) -> Path:
    """
    Plot a normalised confusion matrix as a seaborn heatmap.

    Parameters
    ----------
    preds : np.ndarray, shape (N,)
        Predicted class indices.
    labels : np.ndarray, shape (N,)
        Ground-truth class indices.
    label_names : list[str]
        Ordered class names corresponding to class indices.
    save_dir : str
        Local directory to save the figure into.

    Returns
    -------
    Path
        Local path of the saved figure (confusion_matrix.png).
    """
    cm = confusion_matrix(labels, preds, normalize="true")

    fig, ax = plt.subplots(figsize=(7, 6))
    sns.heatmap(
        cm,
        annot=True,
        fmt=".2f",
        cmap="Blues",
        xticklabels=label_names,
        yticklabels=label_names,
        linewidths=0.5,
        ax=ax,
    )
    ax.set_title("Confusion Matrix (normalised)", fontsize=13, fontweight="bold")
    ax.set_xlabel("Predicted")
    ax.set_ylabel("True")
    fig.tight_layout()

    save_path = Path(save_dir) / "confusion_matrix.png"
    save_path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(save_path, dpi=150, bbox_inches="tight")
    plt.close(fig)

    logger.info("Confusion matrix saved to %s", save_path)
    return save_path


# =============================================================================
# SECTION 5 — Plot: per-class metrics
# =============================================================================
def plot_per_class_metrics(report: dict, save_dir: str) -> Path:
    """
    Plot a grouped bar chart of precision, recall, and F1 per class.

    Parameters
    ----------
    report : dict
        sklearn classification_report output_dict=True — from
        compute_metrics()["report"]. Keys are class names plus
        "accuracy", "macro avg", "weighted avg".
    save_dir : str
        Local directory to save the figure into.

    Returns
    -------
    Path
        Local path of the saved figure (per_class_metrics.png).
    """
    class_names = [k for k, v in report.items() if isinstance(v, dict)]

    precision = [report[c]["precision"] for c in class_names]
    recall = [report[c]["recall"]    for c in class_names]
    f1 = [report[c]["f1-score"]  for c in class_names]

    x = np.arange(len(class_names))
    width  = 0.25

    fig, ax = plt.subplots(figsize=(10, 5))
    ax.bar(x - width, precision, width, label="Precision")
    ax.bar(x, recall,    width, label="Recall")
    ax.bar(x + width, f1, width, label="F1")

    ax.set_title("Per-class Metrics", fontsize=13, fontweight="bold")
    ax.set_xticks(x)
    ax.set_xticklabels(class_names)
    ax.set_ylim(0, 1.05)
    ax.set_ylabel("Score")
    ax.legend()
    ax.grid(axis="y", linestyle="--", alpha=0.5)
    fig.tight_layout()

    save_path = Path(save_dir) / "per_class_metrics.png"
    save_path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(save_path, dpi=150, bbox_inches="tight")
    plt.close(fig)

    logger.info("Per-class metrics saved to %s", save_path)
    return save_path


# =============================================================================
# SECTION 6 — Word document builder
# =============================================================================
def build_word_report(
    cfg: DictConfig,
    run_id: str,
    run_data: dict,
    figure_paths: dict[str, Path],
    save_dir: str,
) -> Path:
    """
    Build a structured Word document embedding all figures.

    Parameters
    ----------
    cfg : DictConfig
        Full Hydra config — reads project.name, model.*, training.*.
    run_id : str
        MLflow run ID — included in the document for traceability.
    run_data : dict
        Output of fetch_run_data() — keys: params, metrics, history.
    figure_paths : dict[str, Path]
        Paths to generated figures. Expected keys:
        training_curves, confusion_matrix, per_class_metrics.
    save_dir : str
        Local directory to save the .docx file into.

    Returns
    -------
    Path
        Local path of the saved Word document (model_report.docx).
    """
    from datetime import datetime, timezone

    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    title = doc.add_heading(cfg.project.name, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph("Model Evaluation Report")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(13)

    ts = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    meta = doc.add_paragraph(f"Generated: {ts}  |  Run ID: {run_id}  |  Environment: {cfg.environment.name}")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.runs[0].font.size = Pt(9)

    doc.add_page_break()

    # 1. Executive Summary
    doc.add_heading("1. Executive Summary", level=1)
    test_acc = run_data["metrics"].get("test_accuracy")
    best_val = run_data["metrics"].get("best_val_acc")
    test_acc_str = f"{test_acc:.4f}" if test_acc is not None else "N/A"
    best_val_str = f"{best_val:.4f}" if best_val is not None else "N/A"
    doc.add_paragraph(
        f"A BERT-based text classifier was fine-tuned on the BBC News dataset to predict "
        f"one of {cfg.model.num_labels} topic categories. "
        f"The model achieved a best validation accuracy of {best_val_str} "
        f"and a test accuracy of {test_acc_str}."
    )

    # 2. Model & Training Configuration
    doc.add_heading("2. Model & Training Configuration", level=1)
    params = run_data["params"]
    table  = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Parameter"
    table.rows[0].cells[1].text = "Value"
    config_rows = [
        ("Base model",              params.get("bert_base_model",         cfg.model.bert_base_model)),
        ("Max sequence length",     params.get("max_seq_length",          cfg.model.max_seq_length)),
        ("Epochs",                  params.get("epochs",                  cfg.training.epochs)),
        ("Batch size",              params.get("batch_size",              cfg.training.batch_size)),
        ("Learning rate",           params.get("lr",                      cfg.training.lr)),
        ("Warmup steps",            params.get("warmup_steps",            cfg.training.warmup_steps)),
        ("Weight decay",            params.get("weight_decay",            cfg.training.weight_decay)),
        ("Early stopping patience", params.get("early_stopping_patience", cfg.training.early_stopping_patience)),
        ("Num labels",              params.get("num_labels",              cfg.model.num_labels)),
    ]
    for name, value in config_rows:
        row = table.add_row()
        row.cells[0].text = name
        row.cells[1].text = str(value)

    # 3. Training Curves
    doc.add_heading("3. Training Curves", level=1)
    doc.add_paragraph(
        "The charts below show training and validation loss and accuracy across epochs. "
        "Convergence of validation loss indicates effective learning without overfitting."
    )
    if "training_curves" in figure_paths:
        doc.add_picture(str(figure_paths["training_curves"]), width=Inches(6.0))

    # 4. Evaluation Metrics
    doc.add_heading("4. Evaluation Metrics", level=1)
    doc.add_paragraph(
        f"Test accuracy: {test_acc_str}. "
        "The chart below shows precision, recall, and F1-score per class."
    )
    if "per_class_metrics" in figure_paths:
        doc.add_picture(str(figure_paths["per_class_metrics"]), width=Inches(6.0))

    # 5. Confusion Matrix
    doc.add_heading("5. Confusion Matrix", level=1)
    doc.add_paragraph(
        "Each cell shows the proportion of true-class samples predicted as each class "
        "(rows normalised to 1). Off-diagonal values indicate misclassifications."
    )
    if "confusion_matrix" in figure_paths:
        doc.add_picture(str(figure_paths["confusion_matrix"]), width=Inches(4.5))

    # 6. Conclusion
    doc.add_heading("6. Conclusion", level=1)
    doc.add_paragraph(
        f"The fine-tuned BERT model (run: {run_id}) demonstrates strong performance on "
        f"BBC News topic classification with a test accuracy of {test_acc_str}. "
        f"Model artefacts and predictions are stored in the "
        f"{cfg.environment.name} GCS bucket for downstream use."
    )

    save_path = Path(save_dir) / "model_report.docx"
    save_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(save_path))

    logger.info("Word report saved to %s", save_path)
    return save_path


# =============================================================================
# SECTION 7 — Upload outputs to GCS
# =============================================================================
def upload_outputs(
    local_dir: str,
    gcs_output_dir: str,
    gcp_project: str,
) -> None:
    """
    Upload every file in local_dir to a GCS directory.

    Parameters
    ----------
    local_dir : str
        Local directory containing files to upload (figures + Word doc).
    gcs_output_dir : str
        GCS URI prefix to upload into, e.g.
        ``gs://bucket/outputs/2026-05-28T10-30-00/``
    gcp_project : str
        GCP project ID for the storage client.
    """
    path_no_scheme = gcs_output_dir.rstrip("/").replace("gs://", "")
    bucket_name    = path_no_scheme.split("/")[0]
    prefix         = "/".join(path_no_scheme.split("/")[1:])

    client = storage.Client(project=gcp_project)
    bucket = client.bucket(bucket_name)

    for file_path in Path(local_dir).iterdir():
        if not file_path.is_file():
            continue
        blob_name = f"{prefix}/{file_path.name}" if prefix else file_path.name
        bucket.blob(blob_name).upload_from_filename(str(file_path))
        logger.info("Uploaded %s -> gs://%s/%s", file_path.name, bucket_name, blob_name)

    logger.info("All outputs uploaded to %s", gcs_output_dir)


# =============================================================================
# SECTION 8 — Report orchestrator
# =============================================================================
def generate_report(
    cfg: DictConfig,
    run_id: str,
    gcs_predictions_uri: str,
) -> str:
    """
    Full report pipeline: download predictions, fetch MLflow run data,
    generate all plots, build Word document, upload outputs to GCS.

    Flow
    ----
    1. Download predictions Parquet from GCS.
    2. Fetch training history + params + metrics from MLflow.
    3. Generate training curves, confusion matrix, per-class metrics plots.
    4. Build Word report embedding all figures.
    5. Upload figures + Word doc to GCS outputs folder.

    Parameters
    ----------
    cfg : DictConfig
        Full Hydra config — reads environment.*, project.name,
        model.*, training.*.
    run_id : str
        MLflow run ID to pull training data from.
    gcs_predictions_uri : str
        GCS URI of the predictions Parquet written by predict().

    Returns
    -------
    str
        GCS URI of the outputs folder containing all report artefacts.
    """
    from datetime import datetime, timezone

    tracking_uri = OmegaConf.select(
        cfg, "environment.mlflow.tracking_uri", default="mlruns"
    )

    _project_root = Path(__file__).resolve().parent.parent.parent
    local_dir     = "/tmp/report" if os.getenv("CLOUD_ML_PROJECT_ID") else str(_project_root / "reports")
    Path(local_dir).mkdir(parents=True, exist_ok=True)

    # Step 1 — Download predictions
    predictions_path = download_predictions(
        gcs_predictions_uri=gcs_predictions_uri,
        gcp_project=cfg.environment.gcp_project,
        local_dir=local_dir,
    )

    table = pq.read_table(str(predictions_path))
    label_names = sorted({col.replace("prob_", "") for col in table.schema.names if col.startswith("prob_")})
    label2idx = {name: idx for idx, name in enumerate(label_names)}

    pred_labels = table.column("predicted").to_pylist()
    preds = np.array([label2idx[p] for p in pred_labels])

    labels = None
    if "true_label" in table.schema.names:
        true_label_list = table.column("true_label").to_pylist()
        labels = np.array([label2idx[l] for l in true_label_list])

    # Step 2 — Fetch MLflow run data
    run_data = fetch_run_data(run_id=run_id, tracking_uri=tracking_uri)

    # Step 3 — Generate plots
    figure_paths: dict[str, Path] = {}

    figure_paths["training_curves"] = plot_training_curves(
        history=run_data["history"],
        save_dir=local_dir,
    )

    if labels is not None:
        figure_paths["confusion_matrix"] = plot_confusion_matrix(
            preds=preds,
            labels=labels,
            label_names=label_names,
            save_dir=local_dir,
        )

        from news_topic_classifier.modeling.predict import compute_metrics
        metrics     = compute_metrics(preds, labels, dict(enumerate(label_names)))
        figure_paths["per_class_metrics"] = plot_per_class_metrics(
            report=metrics["report"],
            save_dir=local_dir,
        )
        run_data["metrics"].setdefault("test_accuracy", metrics["accuracy"])

    # Step 4 — Build Word report
    build_word_report(
        cfg=cfg,
        run_id=run_id,
        run_data=run_data,
        figure_paths=figure_paths,
        save_dir=local_dir,
    )

    # Step 5 — Upload to GCS
    ts = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H-%M-%S")
    gcs_output_dir = f"gs://{cfg.environment.gcs_bucket_artifacts}/outputs/{ts}/"

    upload_outputs(
        local_dir=local_dir,
        gcs_output_dir=gcs_output_dir,
        gcp_project=cfg.environment.gcp_project,
    )

    logger.info("Report complete. Outputs at %s", gcs_output_dir)
    return gcs_output_dir


# =============================================================================
# SECTION 9 — Smoke test
# =============================================================================
if __name__ == "__main__":

    import hydra
    from omegaconf import OmegaConf

    PROJECT_ROOT = Path(__file__).resolve().parent.parent.parent

    @hydra.main(
        config_path=str(PROJECT_ROOT / "conf"),
        config_name="config",
        version_base=None,
    )
    def main(cfg: DictConfig) -> None:

        print("\n" + "=" * 80)
        print("report.py — smoke test")
        print("=" * 80)

        # Redirect MLflow to local SQLite
        cfg = OmegaConf.merge(
            cfg,
            {"environment": {"mlflow": {"tracking_uri": f"sqlite:///{PROJECT_ROOT}/mlflow.db"}}},
        )

        run_id              = OmegaConf.select(cfg, "report.run_id")
        gcs_predictions_uri = OmegaConf.select(cfg, "report.gcs_predictions_uri")

        if not run_id or not gcs_predictions_uri:
            raise ValueError(
                "Pass both overrides:\n"
                "  report.run_id=<mlflow-run-id>\n"
                "  report.gcs_predictions_uri=gs://..."
            )

        print(f"Run ID : {run_id}")
        print(f"Predictions URI : {gcs_predictions_uri}")

        gcs_output_dir = generate_report(
            cfg=cfg,
            run_id=run_id,
            gcs_predictions_uri=gcs_predictions_uri,
        )

        print(f"\nReport outputs at : {gcs_output_dir}")
        print("Smoke test passed.")

    main()
